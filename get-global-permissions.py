import os
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from docx import Document
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s: %(message)s')

# URLs for the Jira instances
source_url = 'https://source.atlassian.net/secure/admin/GlobalPermissions!default.jspa' #$J$Yt*5$
target_url = 'https://target.atlassian.net/secure/admin/GlobalPermissions!default.jspa'

# Initialize WebDriver (Make sure to have the ChromeDriver in your PATH)
driver = webdriver.Chrome()

def manual_login_and_navigate(url):
    driver.get(url)
    print("Please log in to Jira in the browser that opened and navigate to the Global Permissions page.")
    input("Press Enter here in the terminal after you have logged in and navigated to the Global Permissions page...")

def extract_permissions():
    permissions = {}
    table = driver.find_element(By.ID, 'global_perms')
    rows = table.find_elements(By.TAG_NAME, 'tr')[1:]  # Skip the header row
    
    for row in rows:
        cols = row.find_elements(By.TAG_NAME, 'td')
        if len(cols) == 2:
            try:
                permission_name = cols[0].find_element(By.TAG_NAME, 'strong').text.strip()
                groups = []
                for group_elem in cols[1].find_elements(By.TAG_NAME, 'li'):
                    try:
                        group_name = group_elem.find_element(By.TAG_NAME, 'span').text.strip()
                        groups.append(group_name)
                    except Exception as e:
                        logging.error(f"Error extracting group name: {e}")
                permissions[permission_name] = groups
            except Exception as e:
                logging.error(f"Error extracting permission name: {e}")
    
    return permissions

def add_permissions_to_doc(doc, title, permissions):
    doc.add_heading(title, level=1)
    
    for permission, groups in permissions.items():
        doc.add_heading(permission, level=2)
        if groups:
            for group in groups:
                doc.add_paragraph(group, style='ListBullet')
        else:
            doc.add_paragraph("No groups assigned", style='ListBullet')

def compare_permissions(source_permissions, target_permissions):
    missing_permissions = {}
    for permission, groups in source_permissions.items():
        if permission in target_permissions:
            missing_groups = set(groups) - set(target_permissions[permission])
            if missing_groups:
                missing_permissions[permission] = list(missing_groups)
        else:
            missing_permissions[permission] = groups
    return missing_permissions

# Create a new Document
doc = Document()

# Extract and add permissions from the source instance
manual_login_and_navigate(source_url)
source_permissions = extract_permissions()
#add_permissions_to_doc(doc, 'Source Instance Permissions', source_permissions)

# Extract and add permissions from the target instance
manual_login_and_navigate(target_url)
target_permissions = extract_permissions()
#add_permissions_to_doc(doc, 'Target Instance Permissions', target_permissions)

# Compare permissions and add missing groups to the document
missing_permissions = compare_permissions(source_permissions, target_permissions)
add_permissions_to_doc(doc, 'Missing Permissions in Target Instance', missing_permissions)

# Save the document
doc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'jira_global_permissions.docx')
doc.save(doc_path)
logging.info(f"Document saved to {doc_path}")

# Close the WebDriver
driver.quit()
