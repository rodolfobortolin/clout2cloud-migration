import os
import requests
from requests.auth import HTTPBasicAuth
import logging
from collections import defaultdict
from docx import Document
from tqdm import tqdm  # Import the tqdm library for progress bars

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s: %(message)s')

source_config = {
    'email': 'rodolfobortolin@gmail.com',
    'token': '',
    'base_url': 'https://source.atlassian.net'
}

target_config = {
    'email': 'rodolfobortolin@gmail.com',
    'token': '',
    'base_url': 'https://target.atlassian.net'
}

FIELD_TYPE_MAPPING = {
    "option": "Select List (Single Select)",
    "array": "Select List (Multiple Select)",
    "string": "Text",
    "date": "Date",
    "datetime": "Date/Time",
    "number": "Number",
    "sd-servicelevelagreement": "Service Level Agreement (SLA)",
    "project": "Project Picker",
    "option-with-child": "Select List (cascade)",
    # Add other mappings as needed
}

def get_readable_field_type(api_field_type):
    return FIELD_TYPE_MAPPING.get(api_field_type, api_field_type)

def get_group_members(config, group_name, desc='Fetching group members'):
    start_at = 0
    max_results = 50
    members = []

    total = None  # Initially unknown
    progress_desc = f"{desc} ({group_name})"
    with tqdm(total=total, desc=progress_desc, ncols=100) as pbar:
        while True:
            try:
                url = f"{config['base_url']}/rest/api/2/group/member?groupname={group_name}&startAt={start_at}&maxResults={max_results}"
                auth = HTTPBasicAuth(config['email'], config['token'])
                headers = {"Accept": "application/json"}
                
                response = requests.get(url, auth=auth, headers=headers)
                response.raise_for_status()
                data = response.json()
                members.extend(data.get('values', []))
                
                # Update the progress bar
                if total is None:
                    total = data.get('total', len(data.get('values', [])))
                    pbar.total = total
                pbar.update(len(data.get('values', [])))
                
                if data.get('isLast', True):
                    break

                start_at += max_results
            except requests.exceptions.RequestException as e:
                print(f"An error occurred: {e}. Skipping to the next page.")
                start_at += max_results
                continue
            except KeyError as e:
                print(f"Missing expected key: {e}. Skipping to the next page.")
                start_at += max_results
                continue
            except Exception as e:
                print(f"An unexpected error occurred: {e}. Skipping to the next page.")
                start_at += max_results
                continue
        
        # Ensure the progress bar completes fully
        if pbar.n < pbar.total:
            pbar.n = pbar.total
            pbar.last_print_n = pbar.total
        pbar.close()
    
    return members

def get_all_users_by_license(application_roles, config, desc='Fetching users by license'):
    users_by_license = defaultdict(set)
    
    for role in tqdm(application_roles, desc=desc, ncols=100):
        # Combine both groups and defaultGroups
        all_groups = set(role.get('groups', [])) | set(role.get('defaultGroups', []))
        for group in all_groups:
            members = get_group_members(config, group)
            for member in members:
                if 'emailAddress' in member:
                    users_by_license[role['name']].add(member['emailAddress'])
    
    return users_by_license

# Function to get data from Jira instance with progress bar
def get_jira_data(config, endpoint, desc='Fetching data'):
    url = f"{config['base_url']}{endpoint}"
    auth = HTTPBasicAuth(config['email'], config['token'])
    headers = {"Accept": "application/json"}

    # Make the initial request to get the data
    response = requests.get(url, auth=auth, headers=headers)
    response.raise_for_status()
    data = response.json()

    # Determine the number of elements to set the progress bar total
    elements_count = len(data)
    
    with tqdm(total=elements_count, desc=desc, ncols=100) as pbar:
        # Update the progress bar based on the number of elements
        pbar.update(elements_count)
        
    return data

def search_filters(config, desc='Fetching filters'):
    filters = []
    start_at = 0
    max_results = 50
    total = 1
    with tqdm(total=total, desc=desc, ncols=100) as pbar:
        while True:
            try:
                url = f"{config['base_url']}/rest/api/2/filter/search?expand=description,owner,jql,sharePermissions,editPermissions&startAt={start_at}&maxResults={max_results}"
                response = requests.get(url, headers={"Accept": "application/json"}, auth=HTTPBasicAuth(config['email'], config['token']))
                response.raise_for_status()
                data = response.json()
                filters.extend(data.get('values', []))
                start_at += max_results
                if data.get('isLast', True):
                    break
                total = data.get('total', total)
                pbar.total = total
                pbar.update(len(data.get('values', [])))
            except requests.exceptions.RequestException as e:
                print(f"An error occurred: {e}")
                break
            except KeyError as e:
                print(f"Missing expected key: {e}")
                break
            except Exception as e:
                print(f"An unexpected error occurred: {e}")
                break
        pbar.n = total
        pbar.last_print_n = total
        pbar.close()
    return filters

# Function to search dashboards with pagination and progress bar
def search_dashboards(config, desc='Fetching dashboards'):
    dashboards = []
    start_at = 0
    max_results = 50
    total = 1
    with tqdm(total=total, desc=desc, ncols=100) as pbar:
        while True:
            try:
                url = f"{config['base_url']}/rest/api/2/dashboard?startAt={start_at}&maxResults={max_results}"
                response = requests.get(url, headers={"Accept": "application/json"}, auth=HTTPBasicAuth(config['email'], config['token']))
                response.raise_for_status()
                data = response.json()
                dashboards.extend(data.get('dashboards', []))
                start_at += max_results
                if data.get('isLast', True):
                    break
                total = data.get('total', total)
                pbar.total = total
                pbar.update(len(data.get('dashboards', [])))
            except requests.exceptions.RequestException as e:
                print(f"An error occurred: {e}")
                break
            except KeyError as e:
                print(f"Missing expected key: {e}")
                break
            except Exception as e:
                print(f"An unexpected error occurred: {e}")
                break
        pbar.n = total
        pbar.last_print_n = total
        pbar.close()
    return [dashboard for dashboard in dashboards if dashboard['name'] != 'Default dashboard']

# Function to get notification schemes with pagination and progress bar
# Function to get notification schemes with pagination and progress bar
def get_notification_schemes(config, desc='Fetching notification schemes'):
    notification_schemes = []
    start_at = 0
    max_results = 50
    total = 1
    with tqdm(total=total, desc=desc, ncols=100) as pbar:
        while True:
            try:
                url = f"{config['base_url']}/rest/api/2/notificationscheme/project?startAt={start_at}&maxResults={max_results}"
                response = requests.get(url, headers={"Accept": "application/json"}, auth=HTTPBasicAuth(config['email'], config['token']))
                response.raise_for_status()
                data = response.json()
                notification_schemes.extend(data.get('values', []))
                
                # Update the progress bar
                total = data.get('total', total)
                pbar.total = total
                pbar.update(len(data.get('values', [])))
                
                if data.get('isLast', True):
                    break
                
                start_at += max_results
            except requests.exceptions.RequestException as e:
                print(f"An error occurred: {e}")
                start_at += max_results
                continue
            except KeyError as e:
                print(f"Missing expected key: {e}")
                start_at += max_results
                continue
            except Exception as e:
                print(f"An unexpected error occurred: {e}")
                start_at += max_results
                continue
        
        # Ensure the progress bar completes fully
        if pbar.n < pbar.total:
            pbar.n = total
            pbar.last_print_n = total
        pbar.close()
    
    return notification_schemes
  
  # Function to get notification scheme names for a list of schemes
def get_notification_schemes_with_names(config, schemes, desc='Fetching notification scheme names'):
    for scheme in tqdm(schemes, desc=desc, ncols=100):
        scheme['name'] = get_notification_scheme_name(config, scheme['notificationSchemeId'])
    return schemes

# Cache for notification scheme names to avoid multiple API calls
notification_scheme_name_cache = {}
# Function to get notification scheme name
def get_notification_scheme_name(config, scheme_id):
    if scheme_id in notification_scheme_name_cache:
        return notification_scheme_name_cache[scheme_id]

    url = f"{config['base_url']}/rest/api/2/notificationscheme/{scheme_id}"
    response = requests.get(url, headers={"Accept": "application/json"}, auth=HTTPBasicAuth(config['email'], config['token']))
    response.raise_for_status()
    name = response.json().get('name')
    notification_scheme_name_cache[scheme_id] = name
    return name

# Define endpoints for required data
endpoints = {
    'projects': '/rest/api/3/project',
    'priorities': '/rest/api/3/priority',
    'resolutions': '/rest/api/3/resolution',
    'roles': '/rest/api/3/role',
    'issuetypes': '/rest/api/3/issuetype',
    'customfields': '/rest/api/3/field',
    'statuses': '/rest/api/3/status',
}

# Function to get application roles
def get_application_roles(config):
    url = f"{config['base_url']}/rest/api/2/applicationrole"
    auth = HTTPBasicAuth(config['email'], config['token'])
    headers = {"Accept": "application/json"}
    
    with tqdm(total=1, desc="Fetching application roles", ncols=100) as pbar:
        response = requests.get(url, auth=auth, headers=headers)
        pbar.update(1)
        response.raise_for_status()
        return response.json()

# Fetch data from both instances with progress bars
data_source = {key: get_jira_data(source_config, endpoint, f"Fetching {key} from source") for key, endpoint in endpoints.items()}
data_target = {key: get_jira_data(target_config, endpoint, f"Fetching {key} from target") for key, endpoint in endpoints.items()}

# Fetch filters and dashboards with pagination and progress bars
data_source['filters'] = search_filters(source_config, 'Fetching filters from source')
data_target['filters'] = search_filters(target_config, 'Fetching filters from target')
data_source['dashboards'] = search_dashboards(source_config, 'Fetching dashboards from source')
data_target['dashboards'] = search_dashboards(target_config, 'Fetching dashboards from target')

# Fetch notification schemes with pagination and progress bars
notification_schemes_source = get_notification_schemes(source_config, 'Fetching notification schemes from source')
notification_schemes_target = get_notification_schemes(target_config, 'Fetching notification schemes from target')

# Fetch application roles (licenses)
application_roles_source = get_application_roles(source_config)
application_roles_target = get_application_roles(target_config)

# Create a new Document
doc = Document()

def analyze_additions(source, target):
    additions = {}
    for key, value in source.items():
        if key not in target:
            additions[key] = value
    return additions

def analyze_merges(source, target):
    merges = {}
    for key, value in source.items():
        if key in target:
            merges[key] = value
    return merges

# Section functions for each entity
def add_projects_section(doc, source_data, target_data):
    source_count = len(source_data)
    target_count = len(target_data)

    additions = analyze_additions({item['key']: item.get('description', 'No description') for item in source_data}, {item['key']: item.get('description', 'No description') for item in target_data})
    conflicts = analyze_merges({item['key']: item.get('description', 'No description') for item in source_data}, {item['key']: item.get('description', 'No description') for item in target_data})
    
    doc.add_heading('Projects', level=1)
    doc.add_paragraph(f"• Number of projects in source instance: {source_count}")
    doc.add_paragraph(f"• Number of projects in target instance: {target_count}")

    doc.add_heading('Items to be added to the target instance', level=2)
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Key'
        hdr_cells[1].text = 'Description'
        for addition_key, addition_description in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = addition_key
            row_cells[1].text = addition_description
    else:
        doc.add_paragraph("No items identified for addition.")

    if conflicts:
        doc.add_heading('Conflicting project keys requiring renaming for migration', level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Key'
        hdr_cells[1].text = 'Description'
        for conflict_key, conflict_description in conflicts.items():
            row_cells = table.add_row().cells
            row_cells[0].text = conflict_key
            row_cells[1].text = conflict_description

def add_priorities_section(doc, source_data, target_data):
    analyze_and_add_section(doc, 'Priorities', source_data, target_data, 'name')

def add_resolutions_section(doc, source_data, target_data):
    analyze_and_add_section(doc, 'Resolutions', source_data, target_data, 'name')

def add_roles_section(doc, source_data, target_data):
    source_roles = {role['name']: role.get('description', 'No description') for role in source_data}
    target_roles = {role['name']: role.get('description', 'No description') for role in target_data}

    additions = analyze_additions(source_roles, target_roles)
    merges = analyze_merges(source_roles, target_roles)

    doc.add_heading('Project Roles', level=1)
    doc.add_paragraph(f"• Number of project roles in source instance: {len(source_roles)}")
    doc.add_paragraph(f"• Number of project roles in target instance: {len(target_roles)}")

    doc.add_heading('Roles to be added to the target instance', level=2)
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Description'
        for addition_name, addition_description in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = addition_name
            row_cells[1].text = addition_description
    else:
        doc.add_paragraph("No roles identified for addition.")

    doc.add_heading('Roles to be merged due to presence in both instances', level=2)
    if merges:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Description'
        for merge_name, merge_description in merges.items():
            row_cells = table.add_row().cells
            row_cells[0].text = merge_name
            row_cells[1].text = merge_description
    else:
        doc.add_paragraph("No roles identified for merging.")

def add_issuetypes_section(doc, source_data, target_data):
    analyze_and_add_section(doc, 'Issue Types', source_data, target_data, 'name')

def add_filters_section(doc, source_data, target_data):
    source_count = len(source_data)
    target_count = len(target_data)

    doc.add_heading('Filters', level=1)
    doc.add_paragraph(f"• Number of filters in source instance: {source_count}")
    doc.add_paragraph(f"• Number of filters in target instance: {target_count}")

    doc.add_heading('Conflicting filters (same name in both instances)', level=2)
    conflicts = [item['name'] for item in source_data if item['name'] in [i['name'] for i in target_data]]
    if conflicts:
        for conflict in conflicts:
            doc.add_paragraph(f"• {conflict}")
    else:
        doc.add_paragraph("No conflicts identified for filters.")

def add_dashboards_section(doc, source_data, target_data):
    source_count = len(source_data)
    target_count = len(target_data)

    doc.add_heading('Dashboards', level=1)
    doc.add_paragraph(f"• Number of dashboards in source instance: {source_count}")
    doc.add_paragraph(f"• Number of dashboards in target instance: {target_count}")

    doc.add_heading('Conflicting dashboards (same name in both instances)', level=2)
    conflicts = [item['name'] for item in source_data if item['name'] in [i['name'] for i in target_data]]
    if conflicts:
        for conflict in conflicts:
            doc.add_paragraph(f"• {conflict}")
    else:
        doc.add_paragraph("No conflicts identified for dashboards.")

    doc.add_heading('Limitations', level=2)
    doc.add_paragraph("Due to the limitations of the REST API, the following restrictions apply to dashboard migration:")
    doc.add_paragraph("• Unable to migrate dashboard layouts.")
    doc.add_paragraph("• Unable to migrate dashboard ownership. The script will set the user running the script as the owner, and add the original owner as an editor.")
    doc.add_paragraph("• User favorite dashboards will not be retained. Users will need to manually set their favorites in the new instance.")

# Special handling for custom fields to match the required format
def add_custom_fields_section(doc, source_data, target_data):
    doc.add_heading('Custom Fields', level=1)
    
    source_count = len(source_data)
    target_count = len(target_data)
    
    doc.add_paragraph(f"• Number of custom fields in source instance: {source_count}")
    doc.add_paragraph(f"• Number of custom fields in target instance: {target_count}")

    def get_field_type(field):
        return field.get('schema', {}).get('type', 'N/A')

    non_migratable_types = ["option-with-child", "project", "sd-servicelevelagreement", "multiuserpicker"]

    source_fields = {field['name']: get_field_type(field) for field in source_data}
    target_fields = {field['name']: get_field_type(field) for field in target_data}

    # Additions table
    doc.add_heading('Custom fields to be added', level=2)
    additions = {name: source_type for name, source_type in source_fields.items() if name not in target_fields and source_type not in non_migratable_types}
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Source Type'
        for name, source_type in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = get_readable_field_type(source_type)
    else:
        doc.add_paragraph("No custom fields identified for addition.")

    # Merges table
    doc.add_heading('Custom fields with identical names in both instances', level=2)
    merges = {name: source_type for name, source_type in source_fields.items() if name in target_fields and source_type != target_fields[name]}
    if merges:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Source Type'
        hdr_cells[2].text = 'Target Type'
        for name, source_type in merges.items():
            target_type = target_fields.get(name, 'N/A')
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = get_readable_field_type(source_type)
            row_cells[2].text = get_readable_field_type(target_type)
    else:
        doc.add_paragraph("No custom fields with identical names found in both instances with differing types.")

    # Non-migratable fields table
    doc.add_heading('Custom fields that will not be migrated', level=2)
    non_migratable = {name: source_type for name, source_type in source_fields.items() if source_type in non_migratable_types}
    if non_migratable:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Type'
        for name, source_type in non_migratable.items():
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = get_readable_field_type(source_type)
    else:
        doc.add_paragraph("No custom fields identified for exclusion from migration.")

# Special handling for statuses to only include ones with same name but different status categories
def add_statuses_section(doc, source_data, target_data):
    doc.add_heading('Statuses', level=1)

    source_count = len(source_data)
    target_count = len(target_data)
    
    doc.add_paragraph(f"• Number of statuses in source instance: {source_count}")
    doc.add_paragraph(f"• Number of statuses in target instance: {target_count}")

    # Additions table
    doc.add_heading('Statuses to be added', level=2)
    additions = analyze_additions({status['name']: status['statusCategory']['name'] for status in source_data}, {status['name']: status['statusCategory']['name'] for status in target_data})
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Category'
        for addition_name, addition_category in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = addition_name
            row_cells[1].text = addition_category
    else:
        doc.add_paragraph("No statuses identified for addition.")

    # Conflicts table
    doc.add_heading('Statuses with identical names but different categories', level=2)
    source_statuses = {status['name']: status['statusCategory']['name'] for status in source_data}
    target_statuses = {status['name']: status['statusCategory']['name'] for status in target_data}
    conflicts = {name: source_category for name, source_category in source_statuses.items() if name in target_statuses and source_category != target_statuses[name]}
    if conflicts:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Source Category'
        hdr_cells[2].text = 'Target Category'
        hdr_cells[3].text = 'Suggestion'
        for name, source_category in conflicts.items():
            target_category = target_statuses[name]
            suggestion = 'Change category or Merge'
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = source_category
            row_cells[2].text = target_category
            row_cells[3].text = suggestion
    else:
        doc.add_paragraph("No statuses with identical names found but differing categories.")

def add_notification_schemes_section(doc, source_schemes, target_schemes, projects_data, config):
    doc.add_heading('Notification Schemes', level=1)

    # Aggregate projects under the same notification scheme
    scheme_to_projects = defaultdict(list)
    for scheme in source_schemes:
        scheme_name = get_notification_scheme_name(config, scheme['notificationSchemeId'])
        project_name = next((project['name'] for project in projects_data if project['id'] == scheme['projectId']), 'Unknown Project')
        scheme_to_projects[scheme_name].append(project_name)

    source_count = len(scheme_to_projects)
    target_count = len({get_notification_scheme_name(config, scheme['notificationSchemeId']) for scheme in target_schemes})

    doc.add_paragraph(f"• Number of notification schemes in source instance: {source_count}")
    doc.add_paragraph(f"• Number of notification schemes in target instance: {target_count}")

    doc.add_heading('Notification Schemes and Associated Projects', level=2)
    if scheme_to_projects:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Notification Scheme'
        hdr_cells[1].text = 'Project(s)'
        for scheme_name, projects in scheme_to_projects.items():
            row_cells = table.add_row().cells
            row_cells[0].text = scheme_name
            row_cells[1].text = ', '.join(projects)
    else:
        doc.add_paragraph("No notification schemes identified in the source instance.")
    
    doc.add_heading('Migration Note', level=2)
    doc.add_paragraph("All notification schemes from the source instance will be mapped to the default notification scheme in the target instance.")

def add_licenses_section(doc, source_application_roles, target_application_roles, source_config, target_config):
    doc.add_heading('Licenses', level=1)
    
    def add_table(doc, application_roles, heading):
        if application_roles:
            doc.add_heading(heading, level=2)
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Application'
            hdr_cells[1].text = 'Number of Seats'
            hdr_cells[2].text = 'Remaining Seats'
            hdr_cells[3].text = 'User Count'
            hdr_cells[4].text = 'Default Groups'
            hdr_cells[5].text = 'All Groups'
            
            total_seats = 0
            total_remaining_seats = 0
            total_user_count = 0
            
            for role in application_roles:
                total_seats += role['numberOfSeats']
                total_remaining_seats += role['remainingSeats']
                total_user_count += role['userCount']
                
                row_cells = table.add_row().cells
                row_cells[0].text = role['name']
                row_cells[1].text = str(role['numberOfSeats'])
                row_cells[2].text = str(role['remainingSeats'])
                row_cells[3].text = f"{role['userCount']} ({role['userCountDescription']})"
                row_cells[4].text = ', '.join(role['defaultGroups'])
                row_cells[5].text = ', '.join(role['groups'])

            doc.add_paragraph(f"Total number of seats: {total_seats}")
            doc.add_paragraph(f"Total remaining seats: {total_remaining_seats}")
            doc.add_paragraph(f"Total user count: {total_user_count}")
        else:
            doc.add_paragraph(f"No licenses found in {heading.lower()}.")

    add_table(doc, source_application_roles, "Source Instance")
    add_table(doc, target_application_roles, "Target Instance")

    # Get all users by license for source and target instances
    source_users_by_license = get_all_users_by_license(source_application_roles, source_config, 'Fetching users by license (source)')
    target_users_by_license = get_all_users_by_license(target_application_roles, target_config, 'Fetching users by license (target)')

    # Find common users in both instances
    common_users = {}
    for license_type, source_users in source_users_by_license.items():
        if license_type in target_users_by_license:
            target_users = target_users_by_license[license_type]
            common_users[license_type] = source_users & target_users

    # Create table for common users
    doc.add_heading('Common Users and License Savings', level=1)
    if common_users:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'License Type'
        hdr_cells[1].text = 'User'
        hdr_cells[2].text = 'Savings'

        for license_type, users in common_users.items():
            for user in users:
                row_cells = table.add_row().cells
                row_cells[0].text = license_type
                row_cells[1].text = user
                row_cells[2].text = '1 License'  # Since merging will save one license per user

        total_savings = sum(len(users) for users in common_users.values())
        doc.add_paragraph(f"Total Savings: {total_savings} licenses")
    else:
        doc.add_paragraph("No common users found that would save licenses.")
    
    # Mention email domains that will be added (if any)
    doc.add_heading('Domains to be Added', level=1)
    source_domains = set(user.split('@')[-1] for users in source_users_by_license.values() for user in users)
    target_domains = set(user.split('@')[-1] for users in target_users_by_license.values() for user in users)
    new_domains = source_domains - target_domains

    if new_domains:
        doc.add_paragraph("The following domains will be added:")
        for domain in new_domains:
            doc.add_paragraph(domain)
    else:
        doc.add_paragraph("No new domains will be added.")

    # Calculate remaining seats if all users are transferred from source to target
    doc.add_heading('Remaining Seats After Transfer', level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'License Type'
    hdr_cells[1].text = 'Total Users in Source'
    hdr_cells[2].text = 'Total Users in Target'
    hdr_cells[3].text = 'Unique Users from Source'
    hdr_cells[4].text = 'Remaining Seats After Transfer'

    for license_type, source_users in source_users_by_license.items():
        if license_type in target_users_by_license:
            target_users = target_users_by_license[license_type]
            common_users_set = source_users & target_users
            unique_source_users = source_users - common_users_set
            total_users_after_transfer = len(unique_source_users) + len(target_users)
            target_role = next(role for role in target_application_roles if role['name'] == license_type)
            remaining_seats_after_transfer = target_role['numberOfSeats'] - total_users_after_transfer

            row_cells = table.add_row().cells
            row_cells[0].text = license_type
            row_cells[1].text = str(len(source_users))
            row_cells[2].text = str(len(target_users))
            row_cells[3].text = str(len(unique_source_users))
            row_cells[4].text = str(remaining_seats_after_transfer)
        else:
            # Handle cases where the license type exists only in the source
            unique_source_users = source_users
            total_users_after_transfer = len(unique_source_users)
            source_role = next(role for role in source_application_roles if role['name'] == license_type)
            remaining_seats_after_transfer = source_role['numberOfSeats'] - total_users_after_transfer

            row_cells = table.add_row().cells
            row_cells[0].text = license_type
            row_cells[1].text = str(len(source_users))
            row_cells[2].text = '0'
            row_cells[3].text = str(len(unique_source_users))
            row_cells[4].text = str(remaining_seats_after_transfer)

    for license_type, target_users in target_users_by_license.items():
        if license_type not in source_users_by_license:
            row_cells = table.add_row().cells
            row_cells[0].text = license_type
            row_cells[1].text = '0'
            row_cells[2].text = str(len(target_users))
            row_cells[3].text = '0'
            target_role = next(role for role in target_application_roles if role['name'] == license_type)
            remaining_seats_after_transfer = target_role['numberOfSeats'] - len(target_users)
            row_cells[4].text = str(remaining_seats_after_transfer)

# Function to analyze and add sections for all required entities
def analyze_and_add_section(doc, title, source_data, target_data, key_attr):
    source_count = len(source_data)
    target_count = len(target_data)

    additions = analyze_additions({item[key_attr]: item.get('description', 'No description') for item in source_data}, {item[key_attr]: item.get('description', 'No description') for item in target_data})
    merges = analyze_merges({item[key_attr]: item.get('description', 'No description') for item in source_data}, {item[key_attr]: item.get('description', 'No description') for item in target_data})
    
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"• Number of {title.lower()} in source instance: {source_count}")
    doc.add_paragraph(f"• Number of {title.lower()} in target instance: {target_count}")

    doc.add_heading('Items to be added to the target instance', level=2)
    if additions:
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Description'
        for addition_name, addition_description in additions.items():
            row_cells = table.add_row().cells
            row_cells[0].text = addition_name
            row_cells[1].text = addition_description
    else:
        doc.add_paragraph("No items identified for addition.")

    doc.add_heading('Items to be merged due to presence in both instances', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Description'
    for merge_name, merge_description in merges.items():
        row_cells = table.add_row().cells
        row_cells[0].text = merge_name
        row_cells[1].text = merge_description

# Add sections for all required entities with error handling
sections = [
    ('Projects', 'projects', 'key'),
    ('Priorities', 'priorities', 'name'),
    ('Resolutions', 'resolutions', 'name'),
    ('Issue Types', 'issuetypes', 'name'),
    ('Filters', 'filters', 'name'),
    ('Dashboards', 'dashboards', 'name'),
    ('Project Roles', 'roles', 'name')
]

for title, key, attr in sections:
    try:
        if title == 'Projects':
            add_projects_section(doc, data_source[key], data_target[key])
        elif title == 'Filters':
            add_filters_section(doc, data_source[key], data_target[key])
        elif title == 'Dashboards':
            add_dashboards_section(doc, data_source[key], data_target[key])
        elif title == 'Project Roles':
            add_roles_section(doc, data_source[key], data_target[key])
        else:
            analyze_and_add_section(doc, title, data_source[key], data_target[key], attr)
    except Exception as e:
        logging.error(f"Failed to analyze {title}: {e}")

# Special handling for custom fields with error handling
try:
    add_custom_fields_section(doc, data_source['customfields'], data_target['customfields'])
except Exception as e:
    logging.error(f"Failed to analyze custom fields: {e}")

# Special handling for statuses with error handling
try:
    add_statuses_section(doc, data_source['statuses'], data_target['statuses'])
except Exception as e:
    logging.error(f"Failed to analyze statuses: {e}")

# Special handling for notification schemes with error handling
try:
    add_notification_schemes_section(doc, notification_schemes_source, notification_schemes_target, data_source['projects'], source_config)
except Exception as e:
    logging.error(f"Failed to analyze notification schemes: {e}")

# Special handling for licenses with error handling
try:
    add_licenses_section(doc, application_roles_source, application_roles_target, source_config, target_config)
except Exception as e:
    logging.error(f"Failed to analyze licenses: {e}")

# Save the document
doc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'jira_analysis.docx')
doc.save(doc_path)
logging.info(f"Document saved to {doc_path}")
