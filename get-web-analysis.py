import os
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches
import logging
from collections import defaultdict

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s: %(message)s')

# Base URLs for the Jira instances
base_urls = {
    'source': 'https://source.atlassian.net',
    'target': 'https://target.atlassian.net'
}

# Paths for the Jira pages
paths = {
    'global_permissions': '/secure/admin/GlobalPermissions!default.jspa',
    'time_tracking': '/secure/admin/TimeTrackingAdmin.jspa',
    'plans_permissions': '/jira/plans/settings/permissions',
    'issue_hierarchy': '/jira/settings/issues/issue-hierarchy',
    'plans_dependency': '/jira/plans/settings/dependencies',
}

# Initialize WebDriver (Make sure to have the ChromeDriver in your PATH)
driver = webdriver.Chrome()

def manual_login(base_url):
    driver.get(base_url)
    logging.info(f"Please log in to Jira in the browser that opened and navigate to any page of the required Jira instance.")
    input("Press Enter here in the terminal after you have logged in...")

def navigate_to_page(base_url, path):
    url = f"{base_url}{path}"
    driver.get(url)
    time.sleep(4)  # Adjust sleep time as needed to ensure the page loads completely
    logging.info(f"Navigated to {url}")

def extract_permissions():
    permissions = {}
    table = driver.find_element(By.ID, 'global_perms')
    rows = table.find_elements(By.TAG_NAME, 'tr')[1:]  # Skip the header row
    
    for row in rows:
        cols = row.find_elements(By.TAG_NAME, 'td')
        if len(cols) == 2:
            try:
                permission_name = cols[0].find_element(By.TAG_NAME, 'strong').text.strip()
                groups = []
                for group_elem in cols[1].find_elements(By.TAG_NAME, 'li'):
                    try:
                        group_name = group_elem.find_element(By.TAG_NAME, 'span').text.strip()
                        groups.append(group_name)
                    except Exception as e:
                        logging.error(f"Error extracting group name: {e}")
                permissions[permission_name] = groups
            except Exception as e:
                logging.error(f"Error extracting permission name: {e}")
    
    return permissions

def extract_plans_permissions():
    permissions = {}
    table = driver.find_element(By.CSS_SELECTOR, 'table.css-1h2ap37')
    rows = table.find_elements(By.TAG_NAME, 'tr')[1:]  # Skip the header row

    for row in rows:
        cols = row.find_elements(By.TAG_NAME, 'td')
        if len(cols) == 3:
            try:
                permission_name = cols[0].text.strip()
                groups = [group.text.strip() for group in cols[2].find_elements(By.TAG_NAME, 'span')]
                permissions[permission_name] = groups
            except Exception as e:
                logging.error(f"Error extracting plans permission: {e}")

    return permissions

def extract_issue_hierarchy():
    hierarchy = {}
    table = driver.find_element(By.CSS_SELECTOR, 'table[data-testid="admin-pages-issue-hierarchy-directory.ui.table.dynamic-table-unrankable--table"]')
    rows = table.find_elements(By.TAG_NAME, 'tr')[1:]  # Skip the header row
    
    for row in rows:
        cols = row.find_elements(By.TAG_NAME, 'td')
        if len(cols) >= 4:
            try:
                level_name = cols[1].find_element(By.TAG_NAME, 'input').get_attribute('value').strip()
                issue_types = cols[3].text.strip()
                hierarchy[level_name] = issue_types
            except Exception as e:
                logging.error(f"Error extracting issue hierarchy: {e}")
    
    return hierarchy

def extract_plans_details(base_url):
    plans = []
    page = 1
    while True:
        navigate_to_page(base_url, f"/jira/plans?name=&page={page}&sortKey=title&sortOrder=ASC")
        table = driver.find_element(By.CSS_SELECTOR, 'table[aria-label="Plans details"]')
        rows = table.find_elements(By.TAG_NAME, 'tr')[1:]  # Skip the header row

        if not rows:
            break

        for row in rows:
            cols = row.find_elements(By.TAG_NAME, 'td')
            if len(cols) >= 3:
                try:
                    plan_name = cols[1].text.strip()
                    lead = cols[2].text.strip()
                    plans.append((plan_name, lead))
                except Exception as e:
                    logging.error(f"Error extracting plan details: {e}")
        page += 1

    return plans

def add_permissions_to_doc(doc, title, permissions):
    doc.add_heading(title, level=1)
    
    for permission, groups in permissions.items():
        doc.add_heading(permission, level=2)
        if groups:
            for group in groups:
                doc.add_paragraph(group, style='ListBullet')
        else:
            doc.add_paragraph("No groups assigned", style='ListBullet')

def compare_permissions(source_permissions, target_permissions):
    missing_permissions = {}
    for permission, groups in source_permissions.items():
        if permission in target_permissions:
            missing_groups = set(groups) - set(target_permissions[permission])
            if missing_groups:
                missing_permissions[permission] = list(missing_groups)
        else:
            missing_permissions[permission] = groups
    return missing_permissions

def add_issue_hierarchy_to_doc(doc, title, hierarchy):
    doc.add_heading(title, level=1)
    
    for level, issue_types in hierarchy.items():
        doc.add_heading(level, level=2)
        doc.add_paragraph(issue_types, style='ListBullet')

def take_screenshot_of_div(base_url, path, div_identifier, filename, by=By.CLASS_NAME):
    navigate_to_page(base_url, path)
    div_element = driver.find_element(by, div_identifier)
    div_element.screenshot(filename)
    logging.info(f"Screenshot of {div_identifier} saved to {filename}")

def add_screenshot_to_doc(doc, title, filename):
    doc.add_heading(title, level=1)
    doc.add_picture(filename, width=Inches(6))

def compare_plans(source_plans, target_plans):
    source_plan_names = {plan[0] for plan in source_plans}
    target_plan_names = {plan[0] for plan in target_plans}
    
    common_plans = source_plan_names.intersection(target_plan_names)
    return common_plans

def add_plans_details_to_doc(doc, title, plans):
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Plan Name'
    hdr_cells[1].text = 'Lead'
    hdr_cells[2].text = 'Count'

    plan_counts = defaultdict(int)
    for plan_name, lead in plans:
        plan_counts[plan_name] += 1

    for plan_name, count in plan_counts.items():
        lead = next(lead for name, lead in plans if name == plan_name)
        row_cells = table.add_row().cells
        row_cells[0].text = plan_name
        row_cells[1].text = lead
        row_cells[2].text = str(count)

# Create a new Document
doc = Document()

# Define the steps for logging
steps = [
    "Manual login for the source instance",
    "Extract global permissions from source",
    "Manual login for the target instance",
    "Extract global permissions from target",
    "Compare permissions and document",
    "Take screenshots of Time Tracking settings",
    "Extract and compare plans permissions",
    "Take screenshots of issue hierarchy",
    "Take screenshots of plans dependency settings",
    "Extract and document plans details",
    "Save document"
]

# Step-by-step execution with logging
logging.info("Starting process")
try:
    # Step 1: Manual login for the source instance
    logging.info("Step 1: " + steps[0])
    manual_login(base_urls['source'])

    # Step 2: Extract global permissions from the source instance
    logging.info("Step 2: " + steps[1])
    navigate_to_page(base_urls['source'], paths['global_permissions'])
    source_permissions = extract_permissions()

    # Step 3: Manual login for the target instance
    logging.info("Step 3: " + steps[2])
    manual_login(base_urls['target'])

    # Step 4: Extract global permissions from the target instance
    logging.info("Step 4: " + steps[3])
    navigate_to_page(base_urls['target'], paths['global_permissions'])
    target_permissions = extract_permissions()

    # Step 5: Compare permissions and document missing groups
    logging.info("Step 5: " + steps[4])
    missing_permissions = compare_permissions(source_permissions, target_permissions)
    add_permissions_to_doc(doc, 'Missing Permissions in Target Instance', missing_permissions)

    # Step 6: Take screenshots of Time Tracking settings
    logging.info("Step 6: " + steps[5])
    take_screenshot_of_div(base_urls['source'], paths['time_tracking'], 'common-setting-items', 'source_time_tracking.png')
    add_screenshot_to_doc(doc, 'Source Time Tracking Settings', 'source_time_tracking.png')
    take_screenshot_of_div(base_urls['target'], paths['time_tracking'], 'common-setting-items', 'target_time_tracking.png')
    add_screenshot_to_doc(doc, 'Target Time Tracking Settings', 'target_time_tracking.png')

    # Step 7: Extract and compare plans permissions
    logging.info("Step 7: " + steps[6])
    navigate_to_page(base_urls['source'], paths['plans_permissions'])
    source_plans_permissions = extract_plans_permissions()
    navigate_to_page(base_urls['target'], paths['plans_permissions'])
    target_plans_permissions = extract_plans_permissions()
    missing_plans_permissions = compare_permissions(source_plans_permissions, target_plans_permissions)
    add_permissions_to_doc(doc, 'Missing Plans Permissions in Target Instance', missing_plans_permissions)

    # Step 8: Take screenshots of the issue hierarchy page
    logging.info("Step 8: " + steps[7])
    take_screenshot_of_div(base_urls['source'], paths['issue_hierarchy'], 'ak-main-content', 'source_issue_hierarchy.png', By.ID)
    add_screenshot_to_doc(doc, 'Source Issue Hierarchy Settings', 'source_issue_hierarchy.png')
    take_screenshot_of_div(base_urls['target'], paths['issue_hierarchy'], 'ak-main-content', 'target_issue_hierarchy.png', By.ID)
    add_screenshot_to_doc(doc, 'Target Issue Hierarchy Settings', 'target_issue_hierarchy.png')

    # Step 9: Take screenshots of plans dependency settings
    logging.info("Step 9: " + steps[8])
    take_screenshot_of_div(base_urls['source'], paths['plans_dependency'], 'ak-main-content', 'source_plans_dependency.png', By.ID)
    add_screenshot_to_doc(doc, 'Source Plans Dependency Settings', 'source_plans_dependency.png')
    take_screenshot_of_div(base_urls['target'], paths['plans_dependency'], 'ak-main-content', 'target_plans_dependency.png', By.ID)
    add_screenshot_to_doc(doc, 'Target Plans Dependency Settings', 'target_plans_dependency.png')

    # Step 10: Extract and document plans details
    logging.info("Step 10: " + steps[9])
    source_plans = extract_plans_details(base_urls['source'])
    target_plans = extract_plans_details(base_urls['target'])
    
    # Document plans details
    add_plans_details_to_doc(doc, 'Source Plans Details', source_plans)
    add_plans_details_to_doc(doc, 'Target Plans Details', target_plans)
    
    # Check for conflicts
    conflicts = compare_plans(source_plans, target_plans)
    if conflicts:
        doc.add_heading('Plans with Conflicts', level=1)
        for plan_name in conflicts:
            doc.add_paragraph(plan_name, style='ListBullet')

    # Step 11: Save the document
    logging.info("Step 11: " + steps[10])
    doc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'jira_web_analysis.docx')
    doc.save(doc_path)
    logging.info(f"Document saved to {doc_path}")
except Exception as e:
    logging.error(f"An error occurred: {e}")
finally:
    # Close the WebDriver
    driver.quit()
    logging.info("Process completed and WebDriver closed")
