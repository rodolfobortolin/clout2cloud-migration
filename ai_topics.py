import os
import logging
from docx import Document
from openai import OpenAI

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s: %(message)s')

# OpenAI API configuration
OPENAI_API_KEY = ""  # Create one here: https://platform.openai.com/api-keys

def generate_prompt(sections):
    prompt = "I'm creating an assessment for an Atlassian cloud to cloud migration and I need your help with the content to include in the assessment:\n\n"
    
    for section in sections:
        prompt += f"{section['title']}:\n\n"
        for question, answer in section['qa_pairs']:
            prompt += f"{question} {answer}\n"
        prompt += "\n"
    
    prompt += "Please provide detailed content for each of the sections above. Generate the content for each section writing in a professional way. Don't make up new sections or titles and don't include the questions in the assessment. The titles provided are level 2. Which means they have only one '##'"
    return prompt

def get_chatgpt_response(prompt):

    client = OpenAI(api_key=OPENAI_API_KEY)

    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a knowledgeable assistant helping to draft a migration assessment document."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0
        )
        return response.choices[0].message.content     
    
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        return ""

def create_document(doc_path, response_content):
    doc = Document()
    doc.add_heading('Migration Assessment Document', level=1)

    for line in response_content.split('\n'):
        stripped_line = line.strip()
        if stripped_line.startswith('#'):
            level = stripped_line.count('#')
            heading_text = stripped_line.replace('#', '').strip()
            doc.add_heading(heading_text, level=level)
        elif stripped_line:
            if '**' in stripped_line:
                parts = stripped_line.split('**')
                p = doc.add_paragraph()
                for i, part in enumerate(parts):
                    if i % 2 == 0:
                        p.add_run(part)
                    else:
                        p.add_run(part).bold = True
            else:
                doc.add_paragraph(stripped_line)

    doc.save(doc_path)
    logging.info(f"Document saved to {doc_path}")


# Define sections with titles, questions, levels, and user inputs
sections = [
    {
        "title": "Definition of the Migration’s Purpose",
        "level": 1,
        "qa_pairs": [
            ("What are the primary goals of this migration?", input("Enter the primary goals of the migration: ")),
            ("How will the migration support the organization's strategic objectives?", input("Enter how the migration will support the organization's strategic objectives: ")),
            ("What are the expected outcomes and benefits from this migration?", input("Enter the expected outcomes and benefits from the migration: "))
        ]
    },
    {
        "title": "Necessity of the Migration",
        "level": 1,
        "qa_pairs": [
            ("What are the key reasons driving the need for migration?", input("Enter the key reasons driving the need for migration: ")),
            ("Are there any current limitations or issues with the existing instance that necessitate migration?", input("Enter any current limitations or issues with the existing instance that necessitate migration: ")),
            ("How will the migration address these limitations or issues?", input("Enter how the migration will address these limitations or issues: "))
        ]
    },
    {
        "title": "Scope of the Migration",
        "level": 1,
        "qa_pairs": [
            ("Which projects, users, and data will be migrated?", input("Enter which projects, users, and data will be migrated: ")),
            ("Will any third-party plugins or applications be included in the migration? If so, which ones?", input("Enter any third-party plugins or applications to be included in the migration: ")),
            ("What specific projects, users, or data will be excluded from the migration?", input("Enter what specific projects, users, or data will be excluded from the migration: ")),
            ("Will any third-party plugins or applications be excluded from the migration? If so, which ones?", input("Enter any third-party plugins or applications to be excluded from the migration: "))
        ]
    },
    {
        "title": "Deadlines",
        "level": 1,
        "qa_pairs": [
            ("Are there any critical deadlines or milestones that must be met? If yes, explain why.", input("Enter any critical deadlines or milestones that must be met (if yes, explain why): "))
        ]
    }
]

# Generate the prompt and get the response
prompt = generate_prompt(sections)
response_content = get_chatgpt_response(prompt)

# Define the document path
doc_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'migration_assessment.docx')

# Create the document with the responses
create_document(doc_path, response_content)
